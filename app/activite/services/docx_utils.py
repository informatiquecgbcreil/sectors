from __future__ import annotations

import os
import shutil
import subprocess
from dataclasses import dataclass
from datetime import date, datetime
from typing import Any, Iterable

from docx import Document  # fallback python-docx

from app.models import SessionActivite, PresenceActivite, Participant

try:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm
except Exception:  # pragma: no cover
    DocxTemplate = None  # type: ignore
    InlineImage = None  # type: ignore
    Mm = None  # type: ignore


# ---------------------------------------------------------------------
# Configuration / Constants
# ---------------------------------------------------------------------

MONTH_FOLDERS = [
    "01_Janvier",
    "02_Fevrier",
    "03_Mars",
    "04_Avril",
    "05_Mai",
    "06_Juin",
    "07_Juillet",
    "08_Aout",
    "09_Septembre",
    "10_Octobre",
    "11_Novembre",
    "12_Decembre",
]

DEFAULT_SOFFICE_CANDIDATES_WIN = [
    r"C:\Program Files\LibreOffice\program\soffice.com",
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
]

# Env var override (optionnel) : permet de forcer le binaire LibreOffice
# Exemple: LIBREOFFICE_PATH=C:\Program Files\LibreOffice\program\soffice.exe
ENV_LIBREOFFICE_PATH = "LIBREOFFICE_PATH"


# ---------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------

def _safe_filename(s: str) -> str:
    """Filename-safe but human-ish."""
    return "".join(c for c in (s or "") if c.isalnum() or c in ("-", "_", " ")).strip().replace(" ", "_")


def _archives_root(app) -> str:
    root = os.path.join(app.instance_path, "archives_emargements")
    os.makedirs(root, exist_ok=True)
    return root


def _month_folder(month: int) -> str:
    if 1 <= month <= 12:
        return MONTH_FOLDERS[month - 1]
    return f"{month:02d}_Mois"


def _format_date_fr(d: date | None) -> str:
    return d.strftime("%d/%m/%Y") if d else ""


def _find_soffice() -> str | None:
    """
    Locate LibreOffice CLI binary.
    On Windows services, PATH is often incomplete.
    We check:
      1) ENV override (LIBREOFFICE_PATH)
      2) PATH via shutil.which
      3) Common install locations on Windows
    """
    env_path = os.environ.get(ENV_LIBREOFFICE_PATH)
    if env_path and os.path.exists(env_path):
        return env_path

    for cmd in ("soffice.com", "soffice", "libreoffice"):
        p = shutil.which(cmd)
        if p:
            return p

    if os.name == "nt":
        for c in DEFAULT_SOFFICE_CANDIDATES_WIN:
            if os.path.exists(c):
                return c

    return None


def _install_default_templates(app) -> dict[str, str]:
    """
    Ensure default templates exist in instance/docx_templates.
    IMPORTANT: We do NOT overwrite existing instance templates.
    """
    tpl_dir = os.path.join(app.instance_path, "docx_templates")
    os.makedirs(tpl_dir, exist_ok=True)

    assets_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "assets"))

    mapping = {
        "collectif": os.path.join(tpl_dir, "modele_collectif.docx"),
        "individuel": os.path.join(tpl_dir, "modele_individuel.docx"),
    }

    src_collectif = os.path.join(assets_dir, "modele_collectif.docx")
    src_indiv = os.path.join(assets_dir, "modele_individuel.docx")

    if os.path.exists(src_collectif) and not os.path.exists(mapping["collectif"]):
        shutil.copyfile(src_collectif, mapping["collectif"])

    if os.path.exists(src_indiv) and not os.path.exists(mapping["individuel"]):
        shutil.copyfile(src_indiv, mapping["individuel"])

    return mapping


def _month_range(annee: int, mois: int) -> tuple[date, date]:
    """Return [start, end) range for a given month."""
    start = date(annee, mois, 1)
    if mois == 12:
        end = date(annee + 1, 1, 1)
    else:
        end = date(annee, mois + 1, 1)
    return start, end


def _freeze_signature_for_archive(app, out_docx: str, signature_path: str | None, session_id: int | None = None, participant_id: int | None = None) -> str | None:
    """
    Copie la signature dans un sous-dossier 'signatures' à côté du DOCX
    et renvoie le nouveau chemin (stable).
    Ne modifie PAS la DB.

    Objectif: éviter que le DOCX mensuel/collectif dépende d'un dossier tmp,
    d'un instance_path différent (debug vs autre copie), ou d'une purge.
    """
    if not signature_path:
        return None
    if not os.path.exists(signature_path):
        return None

    try:
        out_dir = os.path.dirname(out_docx)
        sig_dir = os.path.join(out_dir, "signatures")
        os.makedirs(sig_dir, exist_ok=True)

        base = os.path.basename(signature_path) or "signature.png"
        name, ext = os.path.splitext(base)
        ext = ext or ".png"

        # Renomme un peu pour éviter collisions et garder lisible
        prefix = "sig"
        if session_id is not None:
            prefix += f"_s{session_id}"
        if participant_id is not None:
            prefix += f"_p{participant_id}"

        safe_name = _safe_filename(name) or "signature"
        target_name = f"{prefix}__{safe_name}{ext}"
        target = os.path.join(sig_dir, target_name)

        if os.path.exists(target):
            return target

        shutil.copy2(signature_path, target)
        return target if os.path.exists(target) else None
    except Exception as e:
        try:
            app.logger.warning("Signature freeze failed: %s (sig=%s out=%s)", e, signature_path, out_docx)
        except Exception:
            pass
        return None


def _docxtpl_inline(template: Any, signature_path: str | None):
    """Return InlineImage (docxtpl) or empty string if missing/unavailable."""
    if not signature_path or not os.path.exists(signature_path):
        return ""
    if InlineImage is None or Mm is None:
        return ""
    try:
        return InlineImage(template, signature_path, width=Mm(30))
    except Exception:
        return ""


# ---------------------------------------------------------------------
# LibreOffice conversion
# ---------------------------------------------------------------------

def _try_docx_to_pdf(app, docx_path: str) -> str | None:
    """
    Convert DOCX to PDF using LibreOffice headless.
    Returns PDF path or None.

    We log failures with useful context instead of swallowing everything silently.
    """
    if not docx_path or not os.path.exists(docx_path):
        return None

    soffice = _find_soffice()
    if not soffice:
        # Useful info for Windows services
        try:
            app.logger.warning("LibreOffice not found (PATH/ENV). Set %s or install LO.", ENV_LIBREOFFICE_PATH)
        except Exception:
            pass
        return None

    out_dir = os.path.dirname(docx_path)
    base = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(out_dir, f"{base}.pdf")

    # Dedicated headless profile (avoid AppData locks under services)
    profile_dir = os.path.join(out_dir, "_lo_profile")
    os.makedirs(profile_dir, exist_ok=True)
    profile_uri = "file:///" + profile_dir.replace("\\", "/")

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nolockcheck",
        "--norestore",
        f"--env:UserInstallation={profile_uri}",
        "--convert-to",
        "pdf",
        "--outdir",
        out_dir,
        docx_path,
    ]

    try:
        res = subprocess.run(
            cmd,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        if os.path.exists(pdf_path):
            return pdf_path

        # LO sometimes generates slightly different casing/extension behaviour; try to locate any pdf with same base
        candidates = [f for f in os.listdir(out_dir) if f.lower() == f"{base}.pdf".lower()]
        if candidates:
            found = os.path.join(out_dir, candidates[0])
            return found if os.path.exists(found) else None

        # Conversion “succeeded” but file missing: log outputs
        try:
            app.logger.warning(
                "LibreOffice conversion reported success but PDF missing. docx=%s stdout=%s stderr=%s",
                docx_path, (res.stdout or "")[:800], (res.stderr or "")[:800]
            )
        except Exception:
            pass
        return None

    except subprocess.CalledProcessError as e:
        try:
            app.logger.error(
                "LibreOffice conversion failed. docx=%s cmd=%s stdout=%s stderr=%s",
                docx_path, cmd, (e.stdout or "")[:800], (e.stderr or "")[:800]
            )
        except Exception:
            pass
        return None
    except Exception as e:
        try:
            app.logger.exception("LibreOffice conversion crashed: %s (docx=%s cmd=%s)", e, docx_path, cmd)
        except Exception:
            pass
        return None


# ---------------------------------------------------------------------
# Generation: COLLECTIF
# ---------------------------------------------------------------------

def generate_collectif_docx_pdf(app, atelier, session: SessionActivite) -> tuple[str, str | None]:
    """
    Generate a DOCX and (if possible) a PDF for a collective session.
    Uses docxtpl if template exists, else python-docx fallback.
    """
    dt = session.date_session or datetime.utcnow().date()
    y, m = dt.year, dt.month

    root = _archives_root(app)
    folder = os.path.join(root, _safe_filename(atelier.secteur), str(y), _safe_filename(atelier.nom), _month_folder(m))
    os.makedirs(folder, exist_ok=True)

    time_label = (session.heure_debut or "")
    if session.heure_fin:
        time_label = f"{time_label}-{session.heure_fin}" if time_label else session.heure_fin

    fname = f"{dt.isoformat()}__COLLECTIF__{_safe_filename(atelier.nom)}__{_safe_filename(time_label or 'session')}.docx"
    out_docx = os.path.join(folder, fname)

    defaults = _install_default_templates(app)
    template_path = atelier.modele_docx_collectif or defaults.get("collectif")

    presences = (
        PresenceActivite.query.filter_by(session_id=session.id)
        .join(Participant)
        .order_by(Participant.nom.asc(), Participant.prenom.asc())
        .all()
    )

    # DOCXTPL path
    if DocxTemplate is not None and template_path and os.path.exists(template_path):
        tpl = DocxTemplate(template_path)
        participants = []
        for pr in presences:
            p = pr.participant

            # Figer la signature dans l'archive du DOCX (stable)
            frozen_sig = _freeze_signature_for_archive(app, out_docx, pr.signature_path, session_id=session.id, participant_id=p.id)

            participants.append(
                {
                    "nom": f"{(p.nom or '').upper()} {(p.prenom or '')}",
                    "email": p.email or "",
                    "ddn": _format_date_fr(p.date_naissance),
                    "sexe": p.genre or "",
                    "type": getattr(p, "type_public", None) or "H",
                    "ville": p.ville or "",
                    "motif": pr.motif or "",
                    "signature": _docxtpl_inline(tpl, frozen_sig),
                }
            )

        context = {
            "lieu": getattr(session, "lieu", None) or "",
            "date": dt.strftime("%d/%m/%Y"),
            "horaires": time_label,
            "titre": atelier.nom,
            "intervenant": "",
            "participants": participants,
        }
        tpl.render(context)
        tpl.save(out_docx)

    # Fallback python-docx
    else:
        doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
        if not doc.paragraphs:
            doc.add_heading(f"Feuille d'émargement - {atelier.nom}", level=1)

        # basic table
        table = doc.tables[0] if doc.tables else doc.add_table(rows=1, cols=8)
        if len(table.rows) == 1:
            headers = ["Nom", "Email", "DDN", "Sexe", "Type", "Ville", "Motif", "Signature"]
            for i, h in enumerate(headers):
                table.cell(0, i).text = h

        # clear existing data rows
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[1]._tr)

        for pr in presences:
            p = pr.participant
            row = table.add_row().cells
            row[0].text = f"{(p.nom or '').upper()} {(p.prenom or '')}"
            row[1].text = p.email or ""
            row[2].text = _format_date_fr(p.date_naissance)
            row[3].text = p.genre or ""
            row[4].text = getattr(p, "type_public", None) or "H"
            row[5].text = p.ville or ""
            row[6].text = pr.motif or ""
            row[7].text = ""
        doc.save(out_docx)

    out_pdf = _try_docx_to_pdf(app, out_docx)
    return out_docx, out_pdf


# ---------------------------------------------------------------------
# Generation: INDIVIDUEL MENSUEL
# ---------------------------------------------------------------------

def generate_individuel_mensuel_docx(app, atelier, annee: int, mois: int) -> str:
    """
    Generate a monthly DOCX for INDIVIDUEL_MENSUEL.
    One row per RDV session (takes the first presence in that session).
    """
    root = _archives_root(app)
    folder = os.path.join(root, _safe_filename(atelier.secteur), str(annee), _safe_filename(atelier.nom), _month_folder(mois))
    os.makedirs(folder, exist_ok=True)

    fname = f"{annee}-{mois:02d}__INDIVIDUEL__{_safe_filename(atelier.nom)}.docx"
    out_docx = os.path.join(folder, fname)

    defaults = _install_default_templates(app)
    template_path = atelier.modele_docx_individuel or defaults.get("individuel")

    start, end = _month_range(annee, mois)

    # Fetch RDV sessions for the month
    sessions = (
        SessionActivite.query
        .filter_by(atelier_id=atelier.id, session_type="INDIVIDUEL_MENSUEL", is_deleted=False)
        .filter(SessionActivite.rdv_date.isnot(None))
        .filter(SessionActivite.rdv_date >= start)
        .filter(SessionActivite.rdv_date < end)
        .all()
    )

    sessions.sort(key=lambda s: (s.rdv_date or start, s.rdv_debut or ""))

    rows: list[dict[str, Any]] = []
    for s in sessions:
        pr = (
            PresenceActivite.query.filter_by(session_id=s.id)
            .join(Participant)
            .order_by(Participant.nom.asc(), Participant.prenom.asc())
            .first()
        )
        if not pr:
            continue

        p = pr.participant
        heures = ""
        if s.rdv_debut and s.rdv_fin:
            heures = f"{s.rdv_debut} - {s.rdv_fin}"
        elif s.rdv_debut:
            heures = s.rdv_debut

        motif = pr.motif or ""
        if getattr(pr, "motif_autre", None):
            motif_autre = pr.motif_autre
            motif = f"{motif} / {motif_autre}" if motif else motif_autre

        # Debug utile: tu verras si les anciens chemins existent vraiment
        try:
            app.logger.info(
                "SIGCHECK: session_id=%s rdv=%s participant=%s sig=%s exists=%s",
                s.id,
                s.rdv_date,
                f"{(p.nom or '').upper()} {(p.prenom or '')}",
                pr.signature_path,
                bool(pr.signature_path and os.path.exists(pr.signature_path)),
            )
        except Exception:
            pass

        # Figer la signature à côté du DOCX mensuel (stable)
        frozen_sig = _freeze_signature_for_archive(app, out_docx, pr.signature_path, session_id=s.id, participant_id=p.id)

        rows.append(
            {
                "nom": f"{(p.nom or '').upper()} {(p.prenom or '')}",
                "email": p.email or "",
                "ddn": _format_date_fr(p.date_naissance),
                "sexe": p.genre or "",
                "type": getattr(p, "type_public", None) or "H",
                "da": _format_date_fr(s.rdv_date),
                "heures": heures,
                "motif": motif,
                "ville": p.ville or "",
                "signature": None,      # docxtpl only
                "_sig_path": frozen_sig,  # chemin figé, pas le tmp
            }
        )

    # Helpful debug logs (remove if you hate logs)
    try:
        app.logger.info("DOCX individuel mensuel: atelier=%s %04d-%02d sessions=%s rows=%s template=%s out=%s",
                        getattr(atelier, "nom", ""), annee, mois, len(sessions), len(rows), template_path, out_docx)
    except Exception:
        pass

    if DocxTemplate is not None and template_path and os.path.exists(template_path):
        tpl = DocxTemplate(template_path)
        for r in rows:
            r["signature"] = _docxtpl_inline(tpl, r.pop("_sig_path", None))

        # IMPORTANT: template should use {{ annee }} (not {{ année }})
        context = {
            "lieu": getattr(atelier, "lieu", None) or "",
            "mois": mois,
            "annee": annee,
            "titre": atelier.nom,
            "intervenant": "",
            "participants": rows,
        }
        tpl.render(context)
        tpl.save(out_docx)

    else:
        # fallback python-docx
        doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
        doc.add_paragraph(f"{atelier.nom} - {mois:02d}/{annee}")

        table = doc.tables[0] if doc.tables else doc.add_table(rows=1, cols=9)
        if len(table.rows) == 1:
            headers = ["Nom", "Email", "DDN", "Sexe", "Type", "Date", "Heures", "Motif", "Ville"]
            for i, h in enumerate(headers):
                table.cell(0, i).text = h

        while len(table.rows) > 1:
            table._tbl.remove(table.rows[1]._tr)

        for r in rows:
            row = table.add_row().cells
            row[0].text = r.get("nom", "")
            row[1].text = r.get("email", "")
            row[2].text = r.get("ddn", "")
            row[3].text = r.get("sexe", "")
            row[4].text = r.get("type", "H")
            row[5].text = r.get("da", "")
            row[6].text = r.get("heures", "")
            row[7].text = r.get("motif", "")
            row[8].text = r.get("ville", "")

        doc.save(out_docx)

    return out_docx


def finalize_individuel_mensuel_pdf(app, atelier, annee: int, mois: int) -> str | None:
    docx_path = generate_individuel_mensuel_docx(app, atelier, annee, mois)
    return _try_docx_to_pdf(app, docx_path)


# ---------------------------------------------------------------------
# Participant bilan
# ---------------------------------------------------------------------

def generate_participant_bilan_docx(app, participant, rows: list[dict]) -> str:
    folder = os.path.join(app.instance_path, "archives_pedagogie")
    os.makedirs(folder, exist_ok=True)

    fname = f"bilan_{participant.id}_{_safe_filename(participant.nom)}_{_safe_filename(participant.prenom)}.docx"
    out_docx = os.path.join(folder, fname)

    template_path = os.path.join(app.instance_path, "docx_templates", "bilan_pedagogique.docx")

    if DocxTemplate is not None and os.path.exists(template_path):
        tpl = DocxTemplate(template_path)
        context = {
            "participant": {
                "nom": participant.nom,
                "prenom": participant.prenom,
                "email": participant.email or "",
                "ville": participant.ville or "",
            },
            "rows": rows,
            "date": date.today().strftime("%d/%m/%Y"),
        }
        tpl.render(context)
        tpl.save(out_docx)

    else:
        doc = Document(template_path) if os.path.exists(template_path) else Document()
        doc.add_heading("Bilan pédagogique", level=1)
        doc.add_paragraph(f"Participant : {participant.nom} {participant.prenom}")
        if participant.email:
            doc.add_paragraph(f"Email : {participant.email}")
        if participant.ville:
            doc.add_paragraph(f"Ville : {participant.ville}")
        doc.add_paragraph(f"Date : {date.today().strftime('%d/%m/%Y')}")

        table = doc.add_table(rows=1, cols=4)
        headers = ["Référentiel", "Compétence", "Date", "Atelier"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        for r in rows:
            row = table.add_row().cells
            row[0].text = r.get("referentiel", "")
            row[1].text = r.get("competence", "")
            row[2].text = r.get("date", "")
            row[3].text = r.get("atelier", "")

        doc.save(out_docx)

    return out_docx


def generate_participant_bilan_pdf(app, participant, rows: list[dict]) -> str | None:
    docx_path = generate_participant_bilan_docx(app, participant, rows)
    return _try_docx_to_pdf(app, docx_path)
